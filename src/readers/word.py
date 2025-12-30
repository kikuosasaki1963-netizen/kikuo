"""Word document reader."""
from __future__ import annotations

from pathlib import Path
from docx import Document


def read_word_file(file_path: str | Path) -> str:
    """Read text content from a Word (.docx) file.

    Args:
        file_path: Path to the Word file.

    Returns:
        Extracted text content.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the file is not a .docx file.
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    if path.suffix.lower() != ".docx":
        raise ValueError(f"Unsupported file format: {path.suffix}. Only .docx is supported.")

    doc = Document(path)
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                paragraphs.append("\t".join(row_text))

    return "\n".join(paragraphs)
